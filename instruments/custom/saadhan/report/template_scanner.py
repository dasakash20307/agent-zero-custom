"""
Template Scanner module for analyzing and extracting templates from existing documents.
"""
from typing import Dict, List, Any, Optional, Union, Type, cast
from typing_extensions import Protocol
from pathlib import Path
import re
import json
from dataclasses import dataclass, asdict
import pandas as pd

from .template import Template, TemplateSection, TemplateBuilder

# Define protocols for document handling
class DocxDocument(Protocol):
    def __call__(self) -> 'DocxDocument': ...
    def add_heading(self, text: str, level: int) -> Any: ...
    def add_paragraph(self, text: str) -> Any: ...
    def add_picture(self, path: str) -> Any: ...
    def save(self, path: str) -> None: ...
    @property
    def paragraphs(self) -> List[Any]: ...
    @property
    def tables(self) -> List[Any]: ...

class PdfReader(Protocol):
    def __init__(self, stream: Any) -> None: ...
    @property
    def pages(self) -> List[Any]: ...

# Type aliases
DocumentType = Type[DocxDocument]
PdfReaderType = Type[PdfReader]

@dataclass
class DocumentStructure:
    """Represents the structure extracted from a document"""
    sections: List[Dict[str, Any]]
    placeholders: List[str]
    format_patterns: Dict[str, List[str]]
    metadata: Dict[str, Any]

class TemplateScanner:
    """Analyzes documents to extract template patterns and structures"""
    
    def __init__(self, template_storage_path: Union[str, Path]):
        self.template_storage_path = Path(template_storage_path)
        self.template_storage_path.mkdir(parents=True, exist_ok=True)
        
        # Common patterns to identify in documents
        self.patterns = {
            'section_headers': r'^[A-Z][^.!?]*[:|\n]',
            'placeholders': r'\{([^}]+)\}',
            'table_headers': r'^[A-Z][a-zA-Z\s]*$',
            'list_items': r'^\s*[-•*]\s',
            'numbered_items': r'^\s*\d+\.',
        }
        
        # Try importing document processing libraries
        try:
            from docx import Document
            self.Document = cast(DocumentType, Document)
        except ImportError:
            self.Document = None
            
        try:
            from PyPDF2 import PdfReader
            self.PdfReader = cast(PdfReaderType, PdfReader)
        except ImportError:
            self.PdfReader = None
    
    def scan_document(self, file_path: Union[str, Path]) -> DocumentStructure:
        """Scan a document and extract its structure"""
        file_path = Path(file_path)
        
        if file_path.suffix.lower() == '.docx':
            if self.Document is None:
                raise ImportError("python-docx package is required for scanning DOCX files")
            return self._scan_docx(file_path)
        elif file_path.suffix.lower() == '.pdf':
            if self.PdfReader is None:
                raise ImportError("PyPDF2 package is required for scanning PDF files")
            return self._scan_pdf(file_path)
        elif file_path.suffix.lower() in ['.txt', '.md']:
            return self._scan_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
    
    def _scan_docx(self, file_path: Path) -> DocumentStructure:
        """Scan a DOCX document"""
        if self.Document is None:
            raise ImportError("python-docx package is required for scanning DOCX files")
            
        try:
            doc = self.Document()
        except Exception as e:
            raise ImportError(f"Failed to create DOCX document: {str(e)}")
            
        sections = []
        placeholders = set()
        format_patterns = {'lists': [], 'tables': [], 'styles': []}
        
        current_section = None
        
        for para in doc.paragraphs:
            # Check if this is a section header
            if re.match(self.patterns['section_headers'], para.text):
                if current_section:
                    sections.append(current_section)
                current_section = {
                    'title': para.text.strip(),
                    'content': [],
                    'style': para.style.name
                }
                format_patterns['styles'].append(para.style.name)
            elif current_section:
                current_section['content'].append(para.text)
                
                # Extract placeholders
                found_placeholders = re.findall(self.patterns['placeholders'], para.text)
                placeholders.update(found_placeholders)
                
                # Check for list items
                if re.match(self.patterns['list_items'], para.text):
                    format_patterns['lists'].append('bullet')
                elif re.match(self.patterns['numbered_items'], para.text):
                    format_patterns['lists'].append('numbered')
        
        if current_section:
            sections.append(current_section)
        
        # Scan tables
        for table in doc.tables:
            format_patterns['tables'].append({
                'rows': len(table.rows),
                'cols': len(table.columns),
                'headers': [cell.text for cell in table.rows[0].cells]
            })
        
        return DocumentStructure(
            sections=sections,
            placeholders=list(placeholders),
            format_patterns=format_patterns,
            metadata={'source': str(file_path), 'type': 'docx'}
        )
    
    def _scan_pdf(self, file_path: Path) -> DocumentStructure:
        """Scan a PDF document"""
        if self.PdfReader is None:
            raise ImportError("PyPDF2 package is required for scanning PDF files")
            
        with open(file_path, 'rb') as file:
            try:
                reader = self.PdfReader(file)
            except Exception as e:
                raise ImportError(f"Failed to create PDF reader: {str(e)}")
                
            sections = []
            placeholders = set()
            format_patterns = {'lists': [], 'tables': [], 'styles': []}
            
            current_section = None
            current_text = ""
            
            for page in reader.pages:
                text = page.extract_text()
                lines = text.split('\n')
                
                for line in lines:
                    # Check if this is a section header
                    if re.match(self.patterns['section_headers'], line):
                        if current_section:
                            current_section['content'] = current_text
                            sections.append(current_section)
                            current_text = ""
                        
                        current_section = {
                            'title': line.strip(),
                            'content': [],
                            'page': reader.pages.index(page) + 1
                        }
                    else:
                        current_text += line + "\n"
                        
                        # Extract placeholders
                        found_placeholders = re.findall(self.patterns['placeholders'], line)
                        placeholders.update(found_placeholders)
                        
                        # Check for list items
                        if re.match(self.patterns['list_items'], line):
                            format_patterns['lists'].append('bullet')
                        elif re.match(self.patterns['numbered_items'], line):
                            format_patterns['lists'].append('numbered')
            
            if current_section:
                current_section['content'] = current_text
                sections.append(current_section)
            
            return DocumentStructure(
                sections=sections,
                placeholders=list(placeholders),
                format_patterns=format_patterns,
                metadata={'source': str(file_path), 'type': 'pdf'}
            )
    
    def _scan_text(self, file_path: Path) -> DocumentStructure:
        """Scan a text document"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            lines = content.split('\n')
            sections = []
            placeholders = set()
            format_patterns = {'lists': [], 'tables': [], 'styles': []}
            
            current_section = None
            current_text = ""
            
            for line in lines:
                # Check if this is a section header
                if re.match(self.patterns['section_headers'], line):
                    if current_section:
                        current_section['content'] = current_text
                        sections.append(current_section)
                        current_text = ""
                    
                    current_section = {
                        'title': line.strip(),
                        'content': []
                    }
                else:
                    current_text += line + "\n"
                    
                    # Extract placeholders
                    found_placeholders = re.findall(self.patterns['placeholders'], line)
                    placeholders.update(found_placeholders)
                    
                    # Check for list items
                    if re.match(self.patterns['list_items'], line):
                        format_patterns['lists'].append('bullet')
                    elif re.match(self.patterns['numbered_items'], line):
                        format_patterns['lists'].append('numbered')
            
            if current_section:
                current_section['content'] = current_text
                sections.append(current_section)
            
            return DocumentStructure(
                sections=sections,
                placeholders=list(placeholders),
                format_patterns=format_patterns,
                metadata={'source': str(file_path), 'type': 'text'}
            )
    
    def create_template(self, doc_structure: DocumentStructure, template_name: str, category: str) -> Template:
        """Create a template from document structure"""
        builder = TemplateBuilder()
        builder.set_name(template_name)
        builder.set_description(f"Template extracted from {doc_structure.metadata['source']}")
        
        # Add sections
        for section in doc_structure.sections:
            content = section.get('content', '')
            if isinstance(content, list):
                content = '\n'.join(content)
                
            metadata = {
                'format_patterns': doc_structure.format_patterns,
                'original_style': section.get('style', '')  # Default to empty string
            }
                
            builder.add_section(
                title=section['title'],
                content_template=content,
                required_fields=doc_structure.placeholders,
                metadata=metadata
            )
        
        template = builder.build()
        self._save_template(template, category)
        return template
    
    def _save_template(self, template: Template, category: str) -> None:
        """Save template to storage"""
        category_path = self.template_storage_path / category
        category_path.mkdir(exist_ok=True)
        
        template_path = category_path / f"{template.name}.json"
        template_data = asdict(template)
        
        with open(template_path, 'w', encoding='utf-8') as f:
            json.dump(template_data, f, indent=2)
    
    def get_similar_templates(self, doc_structure: DocumentStructure, category: str = None) -> List[Template]:
        """Find similar templates based on structure and placeholders"""
        similar_templates = []
        search_path = self.template_storage_path
        
        if category:
            search_path = search_path / category
        
        for template_file in search_path.glob('**/*.json'):
            with open(template_file, 'r', encoding='utf-8') as f:
                template_data = json.load(f)
                template = Template(**template_data)
                
                # Compare placeholders and structure
                similarity_score = self._calculate_similarity(doc_structure, template)
                if similarity_score > 0.6:  # Threshold for similarity
                    similar_templates.append(template)
        
        return similar_templates
    
    def _calculate_similarity(self, doc_structure: DocumentStructure, template: Template) -> float:
        """Calculate similarity between document structure and template"""
        # Compare number of sections
        section_similarity = len(template.sections) / max(len(template.sections), len(doc_structure.sections))
        
        # Compare placeholders
        template_placeholders = set()
        for section in template.sections:
            template_placeholders.update(section.required_fields)
        
        doc_placeholders = set(doc_structure.placeholders)
        placeholder_similarity = len(template_placeholders.intersection(doc_placeholders)) / \
            max(len(template_placeholders), len(doc_placeholders))
        
        # Compare format patterns
        format_similarity = 0
        if hasattr(template.sections[0], 'metadata') and 'format_patterns' in template.sections[0].metadata:
            template_patterns = template.sections[0].metadata['format_patterns']
            doc_patterns = doc_structure.format_patterns
            
            pattern_matches = sum(1 for k in template_patterns if k in doc_patterns)
            format_similarity = pattern_matches / max(len(template_patterns), len(doc_patterns))
        
        # Calculate weighted average
        return (section_similarity * 0.4 + placeholder_similarity * 0.4 + format_similarity * 0.2) 